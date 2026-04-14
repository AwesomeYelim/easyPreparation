"""
문서(docx) 생성 스크립트
usage: python generate_docx.py [출력경로] --data <json_file>

JSON 스키마:
{
  "title":    "문서 제목",
  "subtitle": "부제목",          // optional
  "date":     "2026-03-23",      // optional (기본: 오늘)
  "sections": [
    {
      "heading":     "1. 섹션",  // h1, optional
      "body":        "본문",     // string | list[string], optional
      "code":        "코드 블록", // optional
      "table": {                  // optional
        "headers":    ["컬럼1"],
        "rows":       [["값1"]],
        "col_widths": [8.0],      // optional (cm)
        "header_color": "1A237E"  // optional
      },
      "subsections": [            // optional
        {
          "heading": "1.1 서브섹션",  // h2
          "body":    "...",
          "code":    "...",
          "table":   { ... }
        }
      ]
    }
  ]
}
"""

import sys
import os
import argparse
import json
import tempfile
from datetime import date

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, _HERE)
from design_palette import DOC

# ── 경로 / 상수 ───────────────────────────────────────────────────────────────
_TEMPLATE = os.path.join(os.path.dirname(__file__), "..", "templates", "doc_template.docx")

FONT_BODY      = DOC.FONT_BODY
FONT_CODE      = DOC.FONT_CODE
PAGE_CONTENT_W = 16.0  # cm (A4 21cm - 여백 각 2.5cm)


# ── 문서 초기화 ───────────────────────────────────────────────────────────────

def _load_document():
    if os.path.exists(_TEMPLATE):
        doc = Document(_TEMPLATE)
        body = doc.element.body
        sectPr = body.find(qn("w:sectPr"))
        for child in list(body):
            if child is not sectPr:
                body.remove(child)
    else:
        doc = Document()
    _init_defaults(doc)
    return doc


def _init_defaults(doc):
    normal = doc.styles["Normal"]
    rPr = normal._element.get_or_add_rPr()
    for old in rPr.findall(qn("w:rFonts")):
        rPr.remove(old)
    rFonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), FONT_BODY)
    rPr.insert(0, rFonts)
    for old in rPr.findall(qn("w:sz")):
        rPr.remove(old)
    sz   = OxmlElement("w:sz");   sz.set(qn("w:val"), "20"); rPr.append(sz)
    szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), "20"); rPr.append(szCs)
    pPr = normal._element.get_or_add_pPr()
    for old in pPr.findall(qn("w:spacing")):
        pPr.remove(old)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:after"),    "120")
    sp.set(qn("w:line"),     "280")
    sp.set(qn("w:lineRule"), "exact")
    pPr.append(sp)


# ── 폰트 헬퍼 ─────────────────────────────────────────────────────────────────

def _run_font(run, name=FONT_BODY, size=None, bold=None, color=None, italic=None):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    for old in rPr.findall(qn("w:rFonts")):
        rPr.remove(old)
    rFonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), name)
    rPr.insert(0, rFonts)
    if size   is not None: run.font.size      = size
    if bold   is not None: run.font.bold      = bold
    if color  is not None: run.font.color.rgb = color
    if italic is not None: run.font.italic    = italic


# ── 배경색 ────────────────────────────────────────────────────────────────────

def _cell_bg(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


# ── 테이블 유틸 ───────────────────────────────────────────────────────────────

def _get_or_add_tblPr(tbl):
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    return tblPr


def _table_borders(table, color=DOC.BORDER):
    tblPr = _get_or_add_tblPr(table._tbl)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def _table_cell_margins(table, top=60, left=100, bottom=60, right=100):
    tblPr = _get_or_add_tblPr(table._tbl)
    for old in tblPr.findall(qn("w:tblCellMar")):
        tblPr.remove(old)
    mar = OxmlElement("w:tblCellMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tblPr.append(mar)


def _cm_to_twips(cm: float) -> int:
    return int(cm * 567)


def _set_col_widths(table, widths_cm: list):
    table.autofit = False
    tbl = table._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl.insert(0, tblGrid)
    else:
        for gc in list(tblGrid):
            tblGrid.remove(gc)
    for w in widths_cm:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(_cm_to_twips(w)))
        tblGrid.append(gc)
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            for old in tcPr.findall(qn("w:tcW")):
                tcPr.remove(old)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"),    str(_cm_to_twips(widths_cm[ci])))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)
    tblPr = _get_or_add_tblPr(tbl)
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(_cm_to_twips(sum(widths_cm))))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


# ── 공개 렌더링 헬퍼 ──────────────────────────────────────────────────────────

def styled_table(doc, headers, rows, header_hex=DOC.BG_HEADER, stripe_hex=DOC.BG_STRIPE,
                 col_widths=None, border_color=DOC.BORDER):
    """JSON table 블록 렌더링."""
    n = len(headers)
    table = doc.add_table(rows=1, cols=n)
    try:
        table.style = "Table Normal"
    except KeyError:
        pass
    _table_borders(table, border_color)
    _table_cell_margins(table)
    if col_widths is None:
        w = round(PAGE_CONTENT_W / n, 3)
        col_widths = [w] * n
    _set_col_widths(table, col_widths)

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        _cell_bg(cell, header_hex)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(str(h))
        _run_font(run, size=Pt(9), bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    for ri, row_data in enumerate(rows):
        row = table.add_row()
        bg  = stripe_hex if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ""
            _cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(str(val))
            _run_font(run, size=Pt(9))

    return table


def code_block(doc, text: str):
    """코드 블록 단락 (회색 배경 + 고정폭 폰트)."""
    for line in text.splitlines():
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F5F5F5")
        pPr.append(shd)
        run = p.add_run(line if line else " ")
        _run_font(run, name=FONT_CODE, size=Pt(9))


# ── 섹션 렌더러 ───────────────────────────────────────────────────────────────

def _render_section(doc, sec: dict, level: int = 1):
    """level=1 → Heading 1, level=2 → Heading 2."""
    heading_text = sec.get("heading", "")
    if heading_text:
        p = doc.add_heading(heading_text, level=level)
        colors = {1: RGBColor(*bytes.fromhex(DOC.PRIMARY)), 2: RGBColor(*bytes.fromhex(DOC.SECTION))}
        if p.runs:
            p.runs[0].font.color.rgb = colors.get(level, RGBColor(0, 0, 0))

    body = sec.get("body")
    if body:
        lines = body if isinstance(body, list) else [body]
        for line in lines:
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.size = Pt(10)

    code = sec.get("code")
    if code:
        code_block(doc, code)

    tbl_data = sec.get("table")
    if tbl_data:
        styled_table(
            doc,
            tbl_data["headers"],
            tbl_data["rows"],
            header_hex=tbl_data.get("header_color", DOC.BG_HEADER),
            col_widths=tbl_data.get("col_widths"),
        )

    doc.add_paragraph()

    for sub in sec.get("subsections", []):
        _render_section(doc, sub, level=level + 1)


# ── 빌드 진입점 ───────────────────────────────────────────────────────────────

def build(output_path: str, data: dict) -> str:
    doc = _load_document()

    section = doc.sections[0]
    if section.page_width.cm < 1:
        section.page_width    = Cm(21)
        section.page_height   = Cm(29.7)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    title    = data.get("title",    "문서 제목")
    subtitle = data.get("subtitle", "")
    doc_date = data.get("date",     date.today().strftime("%Y-%m-%d"))

    # 타이틀
    p_title = doc.add_heading(title, level=0)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 부제목 / 날짜
    sub_text = f"{subtitle}  |  {doc_date}" if subtitle else doc_date
    p_sub = doc.add_paragraph(sub_text)
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p_sub.runs:
        p_sub.runs[0].font.color.rgb = RGBColor(*bytes.fromhex(DOC.FG_MUTED))
    doc.add_paragraph()

    for sec in data.get("sections", []):
        _render_section(doc, sec, level=1)

    doc.save(output_path)
    return output_path


# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="docx 생성기 (가이드/매뉴얼)")
    parser.add_argument("output", nargs="?", default=None)
    parser.add_argument("--data", default=None)
    parser.add_argument("--out",  default=None)
    args = parser.parse_args()

    _GUIDE_SCHEMA = {
        "type": "object",
        "required": ["title"],
        "properties": {
            "title":    {"type": "string"},
            "subtitle": {"type": "string"},
            "date":     {"type": "string"},
            "sections": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "heading":     {"type": "string"},
                        "body":        {"oneOf": [{"type": "string"}, {"type": "array"}]},
                        "code":        {"type": "string"},
                        "table":       {"type": "object"},
                        "subsections": {"type": "array"},
                    }
                }
            }
        }
    }

    data = {}
    if args.data:
        data_path = args.data if os.path.isabs(args.data) else os.path.join(os.getcwd(), args.data)
        with open(data_path, encoding="utf-8") as f:
            data = json.load(f)
        try:
            import jsonschema
            jsonschema.validate(instance=data, schema=_GUIDE_SCHEMA)
        except ImportError:
            pass
        except jsonschema.ValidationError as e:
            path = " → ".join(str(p) for p in e.absolute_path)
            print(f"❌ JSON 스키마 오류 [guide]: {e.message}" + (f" (위치: {path})" if path else ""))
            sys.exit(1)

    out_path = args.out or args.output
    if not out_path:
        slug = data.get("title", "document").replace(" ", "_").lower()[:40]
        out_path = os.path.join(
            os.path.dirname(os.path.abspath(__file__)), "..", "output", f"{slug}.docx"
        )

    result = build(out_path, data)
    print(result)


if __name__ == "__main__":
    main()
